
"""




opt = Options()
opt.add_argument(r"--user-data-dir=E:\Hackathon\BrowserChromes\Dunzo")

browser = Chrome(options=opt,service=service)

browser.get("https://www.fao.org/fao-who-codexalimentarius/codex-texts/dbs/pestres/commodities-detail/en/?c_id=131")
time.sleep(3)
Pesticides = browser.find_elements(By.CSS_SELECTOR, 'td[class=" sorting_1"] a')
Limit = browser.find_elements(By.CSS_SELECTOR, 'td[class="mrl_unit"]')
#Pesticides[0].click()
Pest = []
for i in range(len(Pesticides)):
    #print((Pesticides[i].text, Limit[i].text))
    Pest.append((Pesticides[i].text,Limit[i].text.split()[0]))

Sample = browser.find_element(By.CSS_SELECTOR, 'div[class="pestDetail"] h2').text

f = open("%s.txt"%Sample, 'w')
for i in Pest:
    f.write("%s %s\n"%(i[0],i[1]))
    #f.write(*i+"\n")
print(len(Pest))
browser.quit()

import docx
from docx.shared import Pt
def createNew(filename = "Report.docx"):
    doc.add_heading("R.M.D.Engineering College")
    doc.add_paragraph("Lab Details : Rahul Clinic")
    doc.add_heading("Results",1 )
    doc.save(filename)
def ViewTab(textfile):
    count = 0
    f = open("%s"%textfile,"r")
    content = f.read()
    content = content.split("\n")
    sample = content[0]
    s = open("%s.txt"%sample,"r")
    sampleData = s.read().split()
    Pesties = []
    for i in content[1:]:
        Pesties.append(i.split())

    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0].cells
    row[0].text = 'Pesticides Name'
    row[1].text = "PPM or mg/kg"
    row[2].text = "Limit"

    for Pest, ppm in Pesties:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input

        row[0].text = Pest
        row[1].text = ppm
        for i in range (0,len(sampleData),2):
            if sampleData[i] == Pest:
                row[2].text = "<=%s"%sampleData[i+1]
                if ppm>sampleData[i]:
                    count=count+1
    if count == 0:
        doc.add_paragraph("This Food Sample is Edible")
    else:
        doc.add_paragraph("This Food Sample is Non - Edible")
    return "Report.docx"
doc = docx.Document()
createNew()
doc.save(ViewTab("Sample1.txt"))
"""
import time
from selenium.webdriver import Chrome
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys


opt = Options()
opt.add_argument(r"--user-data-dir=E:\Hackathon\BrowserChromes\Dunzo")
service = Service(r"C:\Users\HP\PycharmProjects\WebDriver\chromedriver.exe")
browser = Chrome(options=opt,service=service)
f = open("ConsumerDetails.txt","r")
content = f.read().split("\n")
browser.get("https://borzodelivery.com/in")
time.sleep(3)
data = content[5]
a=browser.find_element(By.CSS_SELECTOR, 'input[placeholder="Pick-up"]')
a.send_keys(data)
time.sleep(1)
a.send_keys(Keys.ARROW_DOWN)
a.send_keys(Keys.ENTER)
a=browser.find_element(By.CSS_SELECTOR, 'input[placeholder="Drop-off"]')
a.send_keys(data)
time.sleep(1)
a.send_keys(Keys.ARROW_DOWN)
a.send_keys(Keys.ENTER)
browser.find_element(By.CSS_SELECTOR,'button[data-testid="unique-mini-form__submit-address"]').click()

browser.find_element(By.CSS_SELECTOR, 'input[data-testid="unique-order-form__point-phone-0"]').send_keys(Keys.ARROW_LEFT*11)
browser.find_element(By.CSS_SELECTOR, 'input[data-testid="unique-order-form__point-phone-0"]').send_keys(content[2])
browser.find_element(By.CSS_SELECTOR, 'textarea[class="OrderTextarea_Root_3bG6h"]').send_keys(content[4])
browser.find_element(By.CSS_SELECTOR, 'input[data-testid="unique-order-form__point-phone-1"]').send_keys(Keys.ARROW_LEFT*11)
browser.find_element(By.CSS_SELECTOR, 'input[data-testid="unique-order-form__point-phone-1"]').send_keys(1212121212)
browser.find_element(By.CSS_SELECTOR, 'textarea[data-testid="unique-order-form__point-note-1"]').send_keys("No4. Annanagar, Limekln st, Chennai")
browser.find_element(By.CSS_SELECTOR, 'li[data-testid="unique-order-form__matter-variant-1"]').click()
browser.find_element(By.CSS_SELECTOR, 'button[data-testid="unique-order-form__create-button"]').click()
