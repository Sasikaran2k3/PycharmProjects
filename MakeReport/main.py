import datetime
import time
import docx
from docx.shared import Cm
from docx.shared import Pt
import docx2pdf
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
def createNew(filename = "Report.docx"):
    head = doc.add_heading()
    head.style.font.size=Pt(18)
    head.add_run(lab).font.underline = True
    doc.add_paragraph("Consumer Name : Rahul ")
    doc.add_paragraph("Age : 20")
    doc.add_paragraph("Reported :%s"%datetime.datetime.now().strftime("%m/%d/%Y, %H:%M:%S"))
    doc.add_heading().add_run("Results:").font.underline = True
    doc.save(filename)
def ViewTab(textfile):
    count = 0
    f = open("%s"%textfile,"r")
    content = f.read()
    content = content.split("\n")
    sample = content[0]
    global lab
    lab = content[1]
    s = open(r"Limits/%s.txt"%sample,"r")
    sampleData = s.read().split()
    Pesties = []

    for i in content[2:]:
        Pesties.append(i.split())

    doc.add_heading("Sample Name : "+sample+"\n",2).style.font.size=Pt(15)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    row = table.rows[0].cells
    row[0].text = "Pesticides Name"
    row[1].text = "PPM or mg/kg"
    row[2].text = "Limit"
    make_rows_bold(table.rows[0])
    for Pest, ppm in Pesties:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input

        row[0].text = Pest
        row[1].text = ppm
        for i in range (0,len(sampleData),2):
            if sampleData[i] == Pest:
                row[2].text = "<=%s"%sampleData[i+1]
                if ppm>sampleData[i+1]:
                    count=count+1
    for row in table.rows:
        row.height = Cm(1.2)
    for col in table.columns:
        col.height = Cm(1.2)
    para = doc.add_paragraph()
    runner = para.add_run("Final Result:")
    runner.font.underline = True
    runner.font.underline = True
    if count == 0:
        res = doc.add_paragraph()
        run = res.add_run("SAFE TO CONSUME").bold = True
        doc.add_paragraph("From this report, we can infer that the food sample is safe to consume as the pesticides present in the sample is within the permitted MRL(Maximum Residue Limit)")
    else:
        res = doc.add_paragraph()
        run = res.add_run("NOT SAFE TO CONSUME").bold = True
        doc.add_paragraph("From this report, we can infer that the food sample is not safe to consume as some of the pesticide present in the sample is above the permitted MRL(Maximum Residue Limit)")
    doc.add_paragraph().add_run("Note:").font.bold = True
    doc.add_paragraph("1. This test is used to detect pesticide in food.\n2.This test is done using liquid chromatography, gas chromatography and uv spectroscopy.")
    doc.add_paragraph().add_run("IMPORTANT INSTRUCTIONS:").font.bold = True
    doc.add_paragraph('*Test results released pertain to the specimen submitted .*All test results are dependent on the quality of the sample received by the Laboratory .*Laboratory investigations are only a tool to facilitate in arriving at a diagnosis and should be clinically correlated by the Referring Physician .*Sample repeats are accepted on request of Referring Physician within 7 days post reporting.*Report delivery may be delayed due to unforeseen circumstances. Inconvenience is regretted.*Certain tests may require further testing at additional cost for derivation of exact value. Kindly submit request within 72 hours post reporting.*The  Courts/Forum  at  Delhi  shall  have  exclusive jurisdiction in all disputes/claims concerning the test(s) & or results of test(s).')
    doc.add_paragraph().add_run("\nYour Truthfully\n(Lab Technician)").font.bold = True
    return r"Reports/Report.docx"
f = open("%s"%"Samples/Sample1.txt","r")
content = f.read()
content = content.split("\n")
print(content)
sample = content[0]
lab = content[1]
doc = docx.Document()
createNew("Reports/Report.docx")
doc.save(ViewTab(r"Samples/Sample1.txt"))
docx2pdf.convert(r"Reports/Report.docx")
f.close()